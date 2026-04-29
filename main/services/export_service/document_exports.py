from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape
from zipfile import ZIP_DEFLATED, ZipFile


def _iter_markdown_blocks(markdown_text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            blocks.append(("blank", ""))
            continue
        if stripped.startswith("# "):
            blocks.append(("h1", stripped[2:].strip()))
        elif stripped.startswith("## "):
            blocks.append(("h2", stripped[3:].strip()))
        elif stripped.startswith("  - "):
            blocks.append(("li2", stripped[4:].strip()))
        elif stripped.startswith("- "):
            blocks.append(("li1", stripped[2:].strip()))
        else:
            blocks.append(("p", stripped))
    return blocks


def _build_fallback_docx_bytes(markdown_text: str, *, run_id: str, generated_at: str, title_text: str) -> bytes:
    body_parts: list[str] = []

    def paragraph_xml(text: str, *, bold: bool = False) -> str:
        escaped = xml_escape(text)
        if bold:
            return (
                '<w:p><w:r><w:rPr><w:b/></w:rPr>'
                f'<w:t xml:space="preserve">{escaped}</w:t></w:r></w:p>'
            )
        return f'<w:p><w:r><w:t xml:space="preserve">{escaped}</w:t></w:r></w:p>'

    body_parts.append(paragraph_xml(title_text, bold=True))
    body_parts.append(paragraph_xml(f"Run ID: {run_id} | Generated: {generated_at}"))

    for kind, text in _iter_markdown_blocks(markdown_text):
        if kind == "blank":
            body_parts.append('<w:p/>')
        elif kind in {"h1", "h2"}:
            body_parts.append(paragraph_xml(text, bold=True))
        elif kind == "li1":
            body_parts.append(paragraph_xml(f"• {text}"))
        elif kind == "li2":
            body_parts.append(paragraph_xml(f"    • {text}"))
        else:
            body_parts.append(paragraph_xml(text))

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
        'xmlns:wne="http://schemas.microsoft.com/office/2006/wordml" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'mc:Ignorable="w14 wp14">'
        '<w:body>' + ''.join(body_parts) +
        '<w:sectPr>'
        '<w:pgSz w:w="12240" w:h="15840"/>'
        '<w:pgMar w:top="1080" w:right="1224" w:bottom="1008" w:left="1224" w:header="708" w:footer="708" w:gutter="0"/>'
        '</w:sectPr>'
        '</w:body></w:document>'
    )
    content_types_xml = """<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>
<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">
  <Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>
  <Default Extension=\"xml\" ContentType=\"application/xml\"/>
  <Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>
</Types>
"""
    rels_xml = """<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>
<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">
  <Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>
</Relationships>
"""
    buffer = BytesIO()
    with ZipFile(buffer, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types_xml)
        archive.writestr("_rels/.rels", rels_xml)
        archive.writestr("word/document.xml", document_xml)
    return buffer.getvalue()


def build_docx_bytes(markdown_text: str, *, run_id: str, generated_at: str, title_text: str = "GridSenpAI Planner Packet") -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

        styles = document.styles
        styles["Normal"].font.name = "Aptos"
        styles["Normal"].font.size = Pt(10.5)
        styles["Normal"].paragraph_format.space_after = Pt(4)

        title = document.add_paragraph()
        title.style = document.styles["Title"]
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run(title_text)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta.paragraph_format.space_after = Pt(10)
        meta_run = meta.add_run(f"Run ID: {run_id} | Generated: {generated_at}")
        meta_run.italic = True
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

        def add_bulleted_paragraph(text: str, level: int = 0) -> None:
            paragraph = document.add_paragraph(style="List Bullet")
            if level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.add_run(text)

        for kind, text in _iter_markdown_blocks(markdown_text):
            if kind == "blank":
                continue
            if kind == "h1":
                heading = document.add_paragraph(style="Heading 1")
                heading.paragraph_format.space_before = Pt(8)
                heading.paragraph_format.space_after = Pt(4)
                run = heading.add_run(text)
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
                continue
            if kind == "h2":
                heading = document.add_paragraph(style="Heading 2")
                heading.paragraph_format.space_before = Pt(10)
                heading.paragraph_format.space_after = Pt(3)
                run = heading.add_run(text)
                run.font.color.rgb = RGBColor(0x26, 0x46, 0x7A)
                continue
            if kind == "li1":
                add_bulleted_paragraph(text, level=0)
                continue
            if kind == "li2":
                add_bulleted_paragraph(text, level=1)
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(text)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(title_text.lower())
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        footer.add_run(" | Page ")

        fld_simple = OxmlElement("w:fldSimple")
        fld_simple.set(qn("w:instr"), "PAGE")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld_simple.append(r)
        footer._element.append(fld_simple)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception:
        return _build_fallback_docx_bytes(
            markdown_text,
            run_id=run_id,
            generated_at=generated_at,
            title_text=title_text,
        )


def build_tldr_docx_bytes(markdown_text: str, *, run_id: str, generated_at: str) -> bytes:
    return build_docx_bytes(
        markdown_text,
        run_id=run_id,
        generated_at=generated_at,
        title_text="GridSenpAI Planner TLDR Summary",
    )


def build_pdf_bytes(markdown_text: str, *, run_id: str, generated_at: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.7 * inch,
        title="GridSenpAI Planner Packet",
    )

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    title_style.textColor = "#1F3A5F"
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=9,
        textColor="#556677",
        spaceAfter=10,
    )
    heading1 = styles["Heading1"]
    heading1.textColor = "#1F3A5F"
    heading1.spaceBefore = 10
    heading1.spaceAfter = 4
    heading2 = styles["Heading2"]
    heading2.textColor = "#26467A"
    heading2.spaceBefore = 10
    heading2.spaceAfter = 4
    body_style = styles["BodyText"]
    body_style.spaceAfter = 4
    bullet_style = ParagraphStyle(
        "Bullet1",
        parent=body_style,
        leftIndent=14,
        bulletIndent=4,
    )
    bullet2_style = ParagraphStyle(
        "Bullet2",
        parent=body_style,
        leftIndent=28,
        bulletIndent=18,
    )

    story: list[Any] = [
        Paragraph("GridSenpAI Planner Packet", title_style),
        Paragraph(f"Run ID: {run_id} | Generated: {generated_at}", subtitle_style),
        Spacer(1, 0.05 * inch),
    ]

    for kind, text in _iter_markdown_blocks(markdown_text):
        if kind == "blank":
            story.append(Spacer(1, 0.05 * inch))
        elif kind == "h1":
            story.append(Paragraph(text, heading1))
        elif kind == "h2":
            story.append(Paragraph(text, heading2))
        elif kind == "li1":
            story.append(Paragraph(text, bullet_style, bulletText="•"))
        elif kind == "li2":
            story.append(Paragraph(text, bullet2_style, bulletText="•"))
        else:
            story.append(Paragraph(text, body_style))

    doc.build(story)
    return buffer.getvalue()


def write_binary(path: str | Path, payload: bytes) -> Path:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_bytes(payload)
    return target
